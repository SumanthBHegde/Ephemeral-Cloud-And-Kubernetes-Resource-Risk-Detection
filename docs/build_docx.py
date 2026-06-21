"""Render docs/report.md into a Societe Generale-styled Word document (docs/report.docx).

Usage:
    python docs/build_docx.py            # reads docs/report.md -> writes docs/report.docx

This is a self-contained, single-dependency (python-docx) renderer. report.md remains the
single source of truth for the report content; this script only formats it. It handles the
controlled Markdown subset the report uses: ATX headings (## / ### / ####), paragraphs with
inline **bold** / *italic* / _italic_ / `code` / [text](url) / <url>, bullet and numbered
lists, GitHub pipe tables, fenced code blocks (```lang), images (![cap](figures/x.png)) and
blockquotes (>). A leading <!--COVER ... --> metadata block drives the cover page.

Design intent: a clean corporate report in the Societe Generale palette (red #E60028 / black /
white) with a cover page, an auto Table of Contents field, page headers/footers with page
numbers, SG-red table headers, shaded monospace code boxes, and numbered figure captions.

Missing image files (e.g. dashboard screenshots not yet captured) degrade gracefully into a
styled placeholder so the build never fails.
"""
from __future__ import annotations

import os
import re
import sys

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is required: pip install python-docx  (or: pip install -r requirements.txt)")

# --- paths ------------------------------------------------------------------------------------
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(DOCS_DIR, "report.md")
OUT_PATH = os.path.join(DOCS_DIR, "report.docx")

# --- Societe Generale theme -------------------------------------------------------------------
SG_RED = RGBColor(0xE6, 0x00, 0x28)
BLACK = RGBColor(0x1D, 0x1D, 0x1B)
GRAY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CODE_FILL = "F4F4F4"          # light-gray code-box / banded-row fill
HEADER_FILL = "E60028"        # SG red table header
BAND_FILL = "F7F2F3"          # very light red tint for banded body rows
CALLOUT_FILL = "FBEAEE"       # light red tint for blockquote callouts

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
CONTENT_WIDTH_IN = 6.3        # usable width inside A4 + ~2.5cm margins

_figure_counter = [0]
_warnings: list[str] = []


# --- low-level docx/XML helpers ---------------------------------------------------------------
def _set_shading(element, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell or paragraph properties element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def _cell_bg(cell, fill_hex: str) -> None:
    _set_shading(cell._tc.get_or_add_tcPr(), fill_hex)


def _para_border(paragraph, *, edge: str, color_hex: str, size: int = 12, space: int = 4) -> None:
    """Add a single border edge ('bottom' | 'left') to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color_hex)
    pBdr.append(el)


def _set_table_borders(table, color_hex: str = "DDDDDD", size: int = 4) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tblPr.append(borders)


def _no_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _add_field(paragraph, instr: str) -> None:
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph run."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve"); instr_el.text = instr
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr_el); run._r.append(end)


# --- inline formatting ------------------------------------------------------------------------
_INLINE_RE = re.compile(
    r"(\*\*.+?\*\*)"          # bold
    r"|(`[^`]+`)"             # code
    r"|(\[[^\]]+\]\([^)]+\))"  # [text](url)
    r"|(<https?://[^>]+>)"    # <url>
    r"|(\*[^*\s][^*]*?\*)"    # *italic*
    r"|(_[^_\s][^_]*?_)"      # _italic_
)


def _add_runs(paragraph, text: str, *, base_color=BLACK, base_size=10.5) -> None:
    """Tokenize a line of inline Markdown and append styled runs to `paragraph`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _plain_run(paragraph, text[pos:m.start()], base_color, base_size)
        tok = m.group(0)
        if tok.startswith("**"):
            r = _plain_run(paragraph, tok[2:-2], base_color, base_size); r.bold = True
        elif tok.startswith("`"):
            r = _plain_run(paragraph, tok[1:-1], BLACK, base_size - 0.5); r.font.name = CODE_FONT
        elif tok.startswith("["):
            label = tok[1:tok.index("]")]
            r = _plain_run(paragraph, label, SG_RED, base_size); r.underline = True
        elif tok.startswith("<"):
            r = _plain_run(paragraph, tok[1:-1], SG_RED, base_size); r.underline = True
        elif tok.startswith("*") or tok.startswith("_"):
            r = _plain_run(paragraph, tok[1:-1], base_color, base_size); r.italic = True
        pos = m.end()
    if pos < len(text):
        _plain_run(paragraph, text[pos:], base_color, base_size)


def _strip_inline(text: str) -> str:
    """Remove inline Markdown markers (**, *, _, `) — used for headings, which carry their own style."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"(?<!\w)[*_]([^*_\s][^*_]*?)[*_](?!\w)", r"\1", text)
    return text


def _plain_run(paragraph, text: str, color, size):
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


# --- block renderers --------------------------------------------------------------------------
def add_heading(doc, text: str, level: int) -> None:
    text = _strip_inline(text)
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    sizes = {1: 15, 2: 12.5, 3: 11}
    colors = {1: SG_RED, 2: BLACK, 3: GRAY}
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, BLACK)
    # tag with a Word heading style so the TOC field can find it
    try:
        p.style = doc.styles[f"Heading {level}"]
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = colors.get(level, BLACK)
        run.bold = True
    except KeyError:
        pass
    if level == 1:
        _para_border(p, edge="bottom", color_hex="E60028", size=12, space=2)


def add_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    _add_runs(p, text)


def add_list_item(doc, text: str, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)


def add_blockquote(doc, lines: list[str]) -> None:
    text = " ".join(l.lstrip("> ").rstrip() for l in lines)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _cell_bg(cell, CALLOUT_FILL)
    _no_table_borders(table)
    cell.width = Inches(CONTENT_WIDTH_IN)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _para_border(p, edge="left", color_hex="E60028", size=18, space=8)
    _add_runs(p, text, base_color=BLACK, base_size=10.5)
    for r in p.runs:
        r.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(CONTENT_WIDTH_IN)
    _cell_bg(cell, CODE_FILL)
    _set_table_borders(table, color_hex="E3E3E3", size=4)
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, alt: str, rel_path: str) -> None:
    _figure_counter[0] += 1
    n = _figure_counter[0]
    abs_path = os.path.normpath(os.path.join(DOCS_DIR, rel_path))
    if os.path.exists(abs_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(abs_path, width=Inches(CONTENT_WIDTH_IN))
        caption_extra = ""
    else:
        _warnings.append(f"missing image: {rel_path}")
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(CONTENT_WIDTH_IN)
        _cell_bg(cell, CODE_FILL)
        _set_table_borders(table, color_hex="E3E3E3", size=4)
        ph = cell.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(18)
        ph.paragraph_format.space_after = Pt(18)
        r = ph.add_run(f"[ figure pending capture — {rel_path} ]")
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = GRAY
        caption_extra = " (pending capture)"
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(f"Figure {n}. {alt}{caption_extra}")
    cr.italic = True
    cr.font.name = BODY_FONT
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = GRAY


def _cell_alignment(spec: str):
    spec = spec.strip()
    if spec.startswith(":") and spec.endswith(":"):
        return WD_ALIGN_PARAGRAPH.CENTER
    if spec.endswith(":"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc, rows: list[str]) -> None:
    header = _split_row(rows[0])
    aligns = [_cell_alignment(c) for c in _split_row(rows[1])]
    body = [_split_row(r) for r in rows[2:]]
    ncols = len(header)

    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table, color_hex="DDDDDD", size=4)

    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        _cell_bg(hdr[i], HEADER_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = aligns[i] if i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        _add_runs(p, text, base_color=WHITE, base_size=9.5)
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = WHITE

    for ri, row in enumerate(body):
        cells = table.add_row().cells
        for i in range(ncols):
            text = row[i] if i < len(row) else ""
            if ri % 2 == 1:
                _cell_bg(cells[i], BAND_FILL)
            p = cells[i].paragraphs[0]
            p.alignment = aligns[i] if i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            _add_runs(p, text, base_color=BLACK, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# --- cover page + TOC -------------------------------------------------------------------------
def parse_cover(md: str) -> dict:
    m = re.search(r"<!--COVER(.*?)-->", md, re.DOTALL)
    meta: dict[str, str] = {}
    if not m:
        return meta
    for line in m.group(1).splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        key, val = line.split(":", 1)
        meta[key.strip()] = val.strip().strip("_").strip()
    return meta


def add_cover(doc, meta: dict) -> None:
    # top SG-red rule
    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(2)
    _para_border(top, edge="bottom", color_hex="E60028", size=24, space=1)
    label = doc.add_paragraph()
    lr = label.add_run((meta.get("subtitle") or "Technical Documentation").upper())
    lr.font.name = BODY_FONT
    lr.font.size = Pt(10)
    lr.bold = True
    lr.font.color.rgb = SG_RED
    label.paragraph_format.space_after = Pt(60)

    if meta.get("project"):
        proj = doc.add_paragraph()
        pr = proj.add_run(meta["project"])
        pr.font.name = BODY_FONT
        pr.font.size = Pt(20)
        pr.bold = True
        pr.font.color.rgb = GRAY
        proj.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph()
    tr = title.add_run(meta.get("title", "Technical Report"))
    tr.font.name = BODY_FONT
    tr.font.size = Pt(28)
    tr.bold = True
    tr.font.color.rgb = BLACK
    title.paragraph_format.space_after = Pt(6)

    if meta.get("track"):
        track = doc.add_paragraph()
        kr = track.add_run(meta["track"])
        kr.font.name = BODY_FONT
        kr.font.size = Pt(11)
        kr.font.color.rgb = SG_RED
        track.paragraph_format.space_after = Pt(40)

    if meta.get("thesis"):
        th = doc.add_paragraph()
        thr = th.add_run("“" + meta["thesis"] + "”")
        thr.italic = True
        thr.font.name = BODY_FONT
        thr.font.size = Pt(12)
        thr.font.color.rgb = BLACK
        th.paragraph_format.left_indent = Inches(0.2)
        _para_border(th, edge="left", color_hex="E60028", size=18, space=10)
        th.paragraph_format.space_after = Pt(60)

    for key, label_text in (("team", "Team"), ("date", "Date"),
                            ("live", "Live demo"), ("video", "Demo video"), ("repo", "Repository")):
        if meta.get(key):
            row = doc.add_paragraph()
            a = row.add_run(f"{label_text}:  ")
            a.bold = True
            a.font.name = BODY_FONT
            a.font.size = Pt(10.5)
            a.font.color.rgb = SG_RED
            b = row.add_run(meta[key])
            b.font.name = BODY_FONT
            b.font.size = Pt(10.5)
            b.font.color.rgb = BLACK
            row.paragraph_format.space_after = Pt(3)

    doc.add_page_break()


def add_toc(doc) -> None:
    h = doc.add_paragraph()
    hr = h.add_run("Contents")
    hr.bold = True
    hr.font.name = BODY_FONT
    hr.font.size = Pt(15)
    hr.font.color.rgb = SG_RED
    _para_border(h, edge="bottom", color_hex="E60028", size=12, space=2)
    h.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose “Update Field” (F9) to build the table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(sep)
    run._r.append(placeholder); run._r.append(end)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = GRAY
        r.font.size = Pt(9)
    doc.add_page_break()


# --- page setup / header / footer -------------------------------------------------------------
def setup_document(doc, meta: dict) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK

    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # header
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"{meta.get('project', 'Sentinel')}  ·  Société Générale Hackathon")
    hr.font.name = BODY_FONT
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY

    # footer: left title, right "Page X of Y"
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run("Ephemeral Cloud & Kubernetes Resource Risk Detection")
    fr.font.name = BODY_FONT
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    tab = fp.add_run("\t\t")
    tab.font.size = Pt(8)
    pg = fp.add_run("Page ")
    pg.font.name = BODY_FONT
    pg.font.size = Pt(8)
    pg.font.color.rgb = GRAY
    _add_field(fp, "PAGE")
    of = fp.add_run(" of ")
    of.font.name = BODY_FONT
    of.font.size = Pt(8)
    of.font.color.rgb = GRAY
    _add_field(fp, "NUMPAGES")


# --- main parse loop --------------------------------------------------------------------------
_SEP_RE = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|[\s:|-]*$")


def _is_table_sep(line: str) -> bool:
    return "|" in line and "-" in line and bool(re.match(r"^[\s:|-]+$", line.strip()))


def render_body(doc, body: str) -> None:
    lines = body.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            i += 1
            continue

        # headings
        m = re.match(r"^(#{2,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1)) - 1  # ## -> 1, ### -> 2, #### -> 3
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # fenced code
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            if lang != "mermaid":   # mermaid is for GitHub; the PNG follows in the report
                add_code_block(doc, buf)
            continue

        # image (whole-line)
        im = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if im:
            add_image(doc, im.group(1), im.group(2))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i])
                i += 1
            add_blockquote(doc, buf)
            continue

        # table: current line has a pipe and the next line is a separator row
        if "|" in line and i + 1 < n and _is_table_sep(lines[i + 1]):
            buf = [lines[i], lines[i + 1]]
            i += 2
            while i < n and "|" in lines[i] and lines[i].strip():
                buf.append(lines[i])
                i += 1
            add_table(doc, buf)
            continue

        # list item
        lm = re.match(r"^\s*([-*])\s+(.*)$", line)
        nm = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if lm:
            add_list_item(doc, lm.group(2).strip(), ordered=False)
            i += 1
            continue
        if nm:
            add_list_item(doc, nm.group(2).strip(), ordered=True)
            i += 1
            continue

        # paragraph: gather consecutive plain lines
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("```")
                    or nxt.startswith(">") or nxt.startswith("![")
                    or re.match(r"^\s*([-*])\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
                    or ("|" in lines[i] and i + 1 < n and _is_table_sep(lines[i + 1]))):
                break
            buf.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(buf))


def main() -> int:
    if not os.path.exists(MD_PATH):
        sys.exit(f"not found: {MD_PATH}")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()

    meta = parse_cover(md)

    # body starts at the first top-level section heading (## ...); everything before
    # (title block + thesis) is represented on the cover page.
    body_start = re.search(r"^##\s", md, re.MULTILINE)
    body = md[body_start.start():] if body_start else md

    doc = Document()
    setup_document(doc, meta)
    add_cover(doc, meta)
    add_toc(doc)
    render_body(doc, body)
    doc.save(OUT_PATH)

    print(f"wrote {OUT_PATH}")
    print(f"figures referenced: {_figure_counter[0]}")
    if _warnings:
        print(f"warnings ({len(_warnings)}):", file=sys.stderr)
        for w in _warnings:
            print(f"  - {w}", file=sys.stderr)
    print("note: open in Word and right-click the Contents field -> Update Field (F9) to build the TOC.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
